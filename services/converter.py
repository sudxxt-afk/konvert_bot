import asyncio
import csv
import hashlib
import json
import os
import re
import shutil
import uuid
import zipfile
from dataclasses import dataclass, replace
from pathlib import Path

from PIL import Image
from pypdf import PdfReader, PdfWriter

try:
    from pillow_heif import register_heif_opener

    register_heif_opener()
except ImportError:
    pass


async def _ffmpeg(*args: str) -> None:
    proc = await asyncio.create_subprocess_exec(
        "ffmpeg", "-y", "-hide_banner", "-loglevel", "error", *args,
        stdout=asyncio.subprocess.DEVNULL,
        stderr=asyncio.subprocess.PIPE,
    )
    _, stderr = await proc.communicate()
    if proc.returncode != 0:
        raise RuntimeError(f"ffmpeg failed: {stderr.decode(errors='ignore')[:500]}")


def _out(src: Path, suffix: str) -> Path:
    return src.with_name(src.stem + suffix)


async def _to_voice(src: Path) -> Path:
    out = _out(src, ".ogg")
    await _ffmpeg("-i", str(src), "-vn", "-ac", "1", "-ar", "32000",
                  "-c:a", "libopus", "-b:a", "32k", str(out))
    return out


async def _to_audio(src: Path, suffix: str, *codec_args: str, limit: tuple | None = None) -> Path:
    out = _out(src, suffix)
    args = ["-i", str(src), "-vn"]
    if limit:
        args += ["-t", str(limit)]
    args += [*codec_args, str(out)]
    await _ffmpeg(*args)
    return out


async def _to_mp4(src: Path) -> Path:
    out = _out(src, ".mp4")
    await _ffmpeg("-i", str(src), "-c:v", "libx264", "-preset", "veryfast",
                  "-crf", "24", "-c:a", "aac", "-b:a", "128k",
                  "-movflags", "+faststart", str(out))
    return out


async def _to_webm(src: Path) -> Path:
    out = _out(src, ".webm")
    await _ffmpeg("-i", str(src), "-c:v", "libvpx-vp9", "-b:v", "0",
                  "-crf", "36", "-deadline", "realtime", "-cpu-used", "5",
                  "-c:a", "libopus", str(out))
    return out


async def _to_gif(src: Path) -> Path:
    out = _out(src, ".gif")
    await _ffmpeg("-i", str(src), "-t", "15",
                  "-vf", "fps=12,scale=480:-2:flags=lanczos", str(out))
    return out


async def _strip_audio(src: Path) -> Path:
    out = _out(src, "_mute.mp4")
    await _ffmpeg("-i", str(src), "-an", "-c:v", "copy", str(out))
    return out


async def _video_sticker(src: Path) -> Path:
    out = _out(src, "_sticker.webm")
    await _ffmpeg(
        "-i", str(src), "-t", "3",
        "-vf", "scale=min(512,iw):-2,fps=30,format=yuv420p",
        "-c:v", "libvpx-vp9", "-b:v", "0", "-crf", "34",
        "-an", str(out),
    )
    return out


async def _gs(*args: str) -> None:
    proc = await asyncio.create_subprocess_exec(
        "gs", "-dNOPAUSE", "-dBATCH", "-dQUIET", *args,
        stdout=asyncio.subprocess.DEVNULL,
        stderr=asyncio.subprocess.PIPE,
    )
    _, stderr = await proc.communicate()
    if proc.returncode != 0:
        raise RuntimeError(f"ghostscript failed: {stderr.decode(errors='ignore')[:500]}")


def _pdf_split_sync(src: Path) -> Path:
    reader = PdfReader(str(src))
    pages = reader.pages
    if len(pages) < 2:
        raise RuntimeError("single page")
    zip_path = _out(src, "_pages.zip")
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for i, page in enumerate(pages, start=1):
            writer = PdfWriter()
            writer.add_page(page)
            part = src.with_name(f"{src.stem}_p{i}.pdf")
            with open(part, "wb") as f:
                writer.write(f)
            zf.write(part, arcname=part.name)
            part.unlink(missing_ok=True)
    return zip_path


async def _pdf_compress(src: Path) -> Path:
    out = _out(src, "_compressed.pdf")
    await _gs(
        "-sDEVICE=pdfwrite",
        "-dCompatibilityLevel=1.4",
        "-dPDFSETTINGS=/ebook",
        "-o", str(out), str(src),
    )
    return out


def _convert_image(src: Path, fmt: str, max_side: int | None = None,
                   quality: int | None = None, grayscale: bool = False,
                   invert: bool = False, rotate: int = 0) -> Path:
    ext = {"JPEG": ".jpg", "PNG": ".png", "WEBP": ".webp", "ICO": ".ico",
           "BMP": ".bmp", "TIFF": ".tiff"}[fmt]
    out = _out(src, ext)
    with Image.open(src) as im:
        im.load()
        if rotate:
            im = im.rotate(rotate, expand=True)
        if grayscale:
            im = im.convert("L")
        elif fmt != "PNG" and im.mode not in ("RGB", "L"):
            im = im.convert("RGB")
        if invert:
            if im.mode == "RGB":
                from PIL import ImageOps
                im = ImageOps.invert(im)
            elif im.mode == "L":
                from PIL import ImageOps
                im = ImageOps.invert(im)
        if fmt == "ICO":
            side = min(im.size)
            im = im.resize((side, side), Image.LANCZOS) if im.size[0] != im.size[1] else im
            im.save(out, format="ICO", sizes=[(side, side)])
            return out
        if max_side and max(im.size) > max_side:
            im.thumbnail((max_side, max_side), Image.LANCZOS)
        kwargs = {}
        if fmt in ("JPEG", "WEBP") and quality:
            kwargs["quality"] = quality
        if fmt == "WEBP":
            kwargs.setdefault("method", 5)
        if fmt == "JPEG" and im.mode == "L":
            pass
        im.save(out, fmt, **kwargs)
    return out


def _sticker(src: Path) -> Path:
    out = _out(src, "_sticker.webp")
    with Image.open(src) as im:
        im.load()
        if im.mode != "RGBA":
            im = im.convert("RGBA")
        im.thumbnail((512, 512), Image.LANCZOS)
        im.save(out, "WEBP", quality=92)
    return out


def eval_fps(expr: str) -> float:
    num, _, den = expr.partition("/")
    try:
        return int(num) / int(den)
    except (ValueError, ZeroDivisionError):
        return 0.0


async def _media_info(src: Path) -> str:
    proc = await asyncio.create_subprocess_exec(
        "ffprobe", "-v", "quiet", "-print_format", "json",
        "-show_format", "-show_streams", str(src),
        stdout=asyncio.subprocess.PIPE, stderr=asyncio.subprocess.DEVNULL,
    )
    stdout, _ = await proc.communicate()
    try:
        data = json.loads(stdout.decode())
    except json.JSONDecodeError:
        raise RuntimeError("ffprobe failed")

    fmt = data.get("format", {})
    lines = []
    name = fmt.get("filename", "")
    if name:
        lines.append(f"📦 {Path(name).suffix.upper().lstrip('.')}")
    size = fmt.get("size")
    if size:
        mb = int(size) / 1024 / 1024
        lines.append(f"⚖️ Размер: {mb:.2f} МБ" if mb >= 1 else f"⚖️ Размер: {int(size)} Б")
    dur = float(fmt.get("duration", 0) or 0)
    if dur:
        m, s = divmod(int(dur), 60)
        lines.append(f"⏱ Длительность: {m}:{s:02d}")
    bitrate = fmt.get("bit_rate")
    if bitrate:
        lines.append(f"📶 Битрейт: {int(bitrate) // 1000} kbps")
    for st in data.get("streams", []):
        ctype = st.get("codec_type")
        if ctype == "video":
            res = st.get("width"), st.get("height")
            fps = st.get("r_frame_rate", "")
            fps_val = round(eval_fps(fps)) if "/" in fps else ""
            lines.append(f"🎬 Видео: {st.get('codec_name')} · {res[0]}×{res[1]}"
                         + (f" · {fps_val} fps" if fps_val else ""))
        elif ctype == "audio":
            ch = st.get("channels", "?")
            sr = st.get("sample_rate", "?")
            lines.append(f"🎧 Аудио: {st.get('codec_name')} · {ch} ch · {sr} Hz")
    if not lines:
        raise RuntimeError("no metadata")
    return "\n".join(lines)


async def _probe_duration(src: Path) -> float:
    proc = await asyncio.create_subprocess_exec(
        "ffprobe", "-v", "quiet", "-print_format", "json",
        "-show_format", str(src),
        stdout=asyncio.subprocess.PIPE, stderr=asyncio.subprocess.DEVNULL,
    )
    stdout, _ = await proc.communicate()
    try:
        data = json.loads(stdout.decode())
        dur = float(data.get("format", {}).get("duration", 0) or 0)
    except (json.JSONDecodeError, ValueError):
        dur = 0
    if dur <= 0:
        raise RuntimeError("no duration")
    return dur


async def _speed(src: Path, factor: float, kind: str) -> Path:
    if kind == VIDEO:
        out = _out(src, "_fast.mp4")
        await _ffmpeg(
            "-i", str(src),
            "-vf", f"setpts=PTS/{factor}",
            "-af", f"atempo={factor}",
            "-c:v", "libx264", "-preset", "veryfast", "-crf", "24",
            "-c:a", "aac", "-b:a", "128k",
            str(out),
        )
    else:
        out = _out(src, "_fast.mp3")
        await _ffmpeg(
            "-i", str(src), "-af", f"atempo={factor}",
            "-c:a", "libmp3lame", "-q:a", "4", str(out),
        )
    return out


async def _trim_clip(src: Path, start: float, duration: float, kind: str) -> Path:
    ss, t = f"{start:.2f}", f"{duration:.2f}"
    if kind in (AUDIO, VOICE):
        out = _out(src, "_cut.mp3")
        await _ffmpeg("-ss", ss, "-t", t, "-i", str(src),
                      "-c:a", "libmp3lame", "-q:a", "4", str(out))
    else:
        out = _out(src, "_cut.mp4")
        await _ffmpeg("-ss", ss, "-t", t, "-i", str(src),
                      "-c:v", "libx264", "-preset", "veryfast", "-crf", "24",
                      "-c:a", "aac", "-b:a", "128k", str(out))
    return out


def _unzip_files(src: Path, dest: Path) -> tuple[list[Path], int]:
    files: list[Path] = []
    skipped = 0
    with zipfile.ZipFile(src) as zf:
        total = sum(1 for i in zf.infolist() if not i.is_dir())
        for info in zf.infolist():
            if info.is_dir():
                continue
            name = info.filename
            parts = [p for p in Path(name).parts if p not in ("/", "\\", ".", "..")]
            if not parts or name.startswith("__MACOSX") or parts[-1].startswith("."):
                continue
            if len(files) >= 20:
                break
            target = dest.joinpath(*parts[-3:])
            target.parent.mkdir(parents=True, exist_ok=True)
            with zf.open(info) as s, open(target, "wb") as f:
                shutil.copyfileobj(s, f)
            files.append(target)
        skipped = max(0, total - len(files))
    return files, skipped


def _hash_file_sync(src: Path) -> str:
    md5 = hashlib.md5()
    sha = hashlib.sha256()
    with open(src, "rb") as f:
        while chunk := f.read(1 << 20):
            md5.update(chunk)
            sha.update(chunk)
    return (
        "#️⃣ <b>Хеши файла</b>\n\n"
        f"<code>MD5</code>\n<code>{md5.hexdigest()}</code>\n\n"
        f"<code>SHA256</code>\n<code>{sha.hexdigest()}</code>"
    )


async def _hash_file(src: Path) -> str:
    return await asyncio.to_thread(_hash_file_sync, src)


FONT_CANDIDATES = [
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
]


def _txt_to_pdf_sync(src: Path) -> Path:
    from fpdf import FPDF

    out = _out(src, ".pdf")
    text = src.read_text(encoding="utf-8", errors="replace")
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    font = next((p for p in FONT_CANDIDATES if Path(p).exists()), None)
    if font:
        pdf.add_font("uni", "", font)
        pdf.set_font("uni", size=11)
    else:
        pdf.set_font("helvetica", size=11)
    for para in text.split("\n"):
        pdf.multi_cell(0, 6, para if para.strip() else "")
    pdf.output(str(out))
    return out


def _docx_to_txt_sync(src: Path) -> Path:
    from docx import Document

    out = _out(src, ".txt")
    doc = Document(str(src))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(c.text for c in row.cells))
    out.write_text("\n".join(parts), encoding="utf-8")
    return out


def _xlsx_to_csv_sync(src: Path) -> Path:
    from openpyxl import load_workbook

    out = _out(src, ".csv")
    wb = load_workbook(str(src), read_only=True, data_only=True)
    ws = wb.active
    with open(out, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        for row in ws.iter_rows(values_only=True):
            writer.writerow(["" if v is None else v for v in row])
    wb.close()
    return out


def _csv_to_xlsx_sync(src: Path) -> Path:
    from openpyxl import Workbook

    out = _out(src, ".xlsx")
    wb = Workbook()
    ws = wb.active
    with open(src, encoding="utf-8-sig", newline="") as f:
        for row in csv.reader(f):
            ws.append(row)
    wb.save(str(out))
    return out


_TS_RE = re.compile(r"^\d{1,2}:\d{2}:\d{2}[,.]\d{1,3}\s*-->")


def _sub_to_txt_sync(src: Path) -> Path:
    out = _out(src, ".txt")
    lines_out = []
    for line in src.read_text(encoding="utf-8", errors="replace").splitlines():
        stripped = line.strip()
        if not stripped or stripped.upper().startswith(("WEBVTT", "NOTE")):
            continue
        if stripped.isdigit():
            continue
        if "-->" in stripped or _TS_RE.match(stripped):
            continue
        lines_out.append(stripped)
    seen_prev = None
    deduped = []
    for l in lines_out:
        if l != seen_prev:
            deduped.append(l)
        seen_prev = l
    out.write_text("\n".join(deduped), encoding="utf-8")
    return out


def _qr_from_text_sync(src: Path) -> Path:
    import qrcode

    out = _out(src, ".png")
    text = src.read_text(encoding="utf-8", errors="replace")[:2000]
    img = qrcode.make(text, box_size=10, border=2)
    img.save(out, "PNG")
    return out


def make_qr_png(text: str) -> Path:
    import qrcode

    out = Path(uuid.uuid4().hex[:16]).with_suffix(".png")
    img = qrcode.make(text, box_size=10, border=2)
    img.save(out, "PNG")
    return out


async def _qr_decode(src: Path) -> str:
    import cv2

    def decode() -> str:
        img = cv2.imread(str(src))
        if img is None:
            raise RuntimeError("not an image")
        detector = cv2.QRCodeDetector()
        data, points, _ = detector.detectAndDecode(img)
        return data

    data = await asyncio.to_thread(decode)
    if not data:
        raise RuntimeError("no qr found")
    return f"🔍 <b>QR распознан</b>\n\n<code>{data}</code>"


@dataclass(frozen=True)
class ActionSpec:
    code: str
    label: str
    kinds: frozenset
    send: str
    run: object
    produces: str = "file"
    group: str = ""


PHOTO = "photo"
VIDEO = "video"
VIDEO_NOTE = "video_note"
AUDIO = "audio"
VOICE = "voice"
PDF = "pdf"
DOCX = "docx"
XLSX = "xlsx"
CSV = "csv"
TEXT = "text"
SUB = "sub"
ZIP = "zip"

MEDIA_KINDS = frozenset({VIDEO, VIDEO_NOTE, AUDIO, VOICE})


def _media_send(kind: str) -> str:
    return "animation" if kind == VIDEO else "audio"

ACTIONS: dict[str, ActionSpec] = {
    spec.code: spec
    for spec in (
        ActionSpec("voice", "🎤 В голосовое", MEDIA_KINDS, "voice",
                   lambda s: _to_voice(s)),
        ActionSpec("mp3", "🎵 MP3", MEDIA_KINDS, "audio",
                   lambda s: _to_audio(s, ".mp3", "-c:a", "libmp3lame", "-q:a", "4")),
        ActionSpec("wav", "🎼 WAV", MEDIA_KINDS, "audio",
                   lambda s: _to_audio(s, ".wav", "-c:a", "pcm_s16le")),
        ActionSpec("m4a", "🎧 M4A", MEDIA_KINDS, "audio",
                   lambda s: _to_audio(s, ".m4a", "-c:a", "aac", "-b:a", "128k")),
        ActionSpec("flac", "💿 FLAC", MEDIA_KINDS, "document",
                   lambda s: _to_audio(s, ".flac", "-c:a", "flac")),
        ActionSpec("ogg", "🌀 OGG", frozenset({AUDIO, VOICE}), "audio",
                   lambda s: _to_audio(s, ".ogg", "-c:a", "libvorbis", "-q:a", "5")),
        ActionSpec("ringtone", "📱 Рингтон M4R", frozenset({AUDIO, VOICE}), "document",
                   lambda s: _to_audio(s, ".m4r", "-c:a", "aac", "-b:a", "192k", limit=30)),
        ActionSpec("gif", "🎞 GIF 15 сек", frozenset({VIDEO}), "animation", lambda s: _to_gif(s)),
        ActionSpec("mp4", "📦 Нормализовать MP4", frozenset({VIDEO}), "animation", lambda s: _to_mp4(s)),
        ActionSpec("webm", "🌐 Сжать в WebM", frozenset({VIDEO}), "animation", lambda s: _to_webm(s)),
        ActionSpec("mute", "🔇 Убрать звук", frozenset({VIDEO}), "animation", lambda s: _strip_audio(s)),
        ActionSpec("vsticker", "🩹 Видео-стикер 3 сек", frozenset({VIDEO}), "animation", lambda s: _video_sticker(s)),
        ActionSpec("speed15", "⚡️ ×1.5 быстрее", MEDIA_KINDS, "auto",
                   lambda s, k: _speed(s, 1.5, k)),
        ActionSpec("speed2", "⚡️ ×2 быстрее", MEDIA_KINDS, "auto",
                   lambda s, k: _speed(s, 2.0, k)),
        ActionSpec("jpg", "🖼 JPEG", frozenset({PHOTO}), "photo",
                   lambda s: asyncio.to_thread(_convert_image, s, "JPEG", quality=90)),
        ActionSpec("png", "🎨 PNG", frozenset({PHOTO}), "photo",
                   lambda s: asyncio.to_thread(_convert_image, s, "PNG")),
        ActionSpec("webp", "🌐 WebP", frozenset({PHOTO}), "document",
                   lambda s: asyncio.to_thread(_convert_image, s, "WEBP", quality=88)),
        ActionSpec("compress", "📦 Сжать фото", frozenset({PHOTO}), "document",
                   lambda s: asyncio.to_thread(_convert_image, s, "JPEG", max_side=1600, quality=60)),
        ActionSpec("sticker", "🩹 Под стикер 512", frozenset({PHOTO}), "document",
                   lambda s: asyncio.to_thread(_sticker, s)),
        ActionSpec("ico", "🎯 Иконка ICO", frozenset({PHOTO}), "document",
                   lambda s: asyncio.to_thread(_convert_image, s, "ICO")),
        ActionSpec("bmp", "🖌 BMP", frozenset({PHOTO}), "document",
                   lambda s: asyncio.to_thread(_convert_image, s, "BMP")),
        ActionSpec("tiff", "🗂 TIFF", frozenset({PHOTO}), "document",
                   lambda s: asyncio.to_thread(_convert_image, s, "TIFF")),
        ActionSpec("gray", "⚫️ Ч/Б", frozenset({PHOTO}), "photo",
                   lambda s: asyncio.to_thread(_convert_image, s, "JPEG", grayscale=True, quality=90)),
        ActionSpec("invert", "🌗 Инверсия", frozenset({PHOTO}), "photo",
                   lambda s: asyncio.to_thread(_convert_image, s, "PNG", invert=True)),
        ActionSpec("rotate", "↻ Повернуть 90°", frozenset({PHOTO}), "photo",
                   lambda s: asyncio.to_thread(_convert_image, s, "JPEG", rotate=-90, quality=92)),
        ActionSpec("qrscan", "🔍 Найти QR", frozenset({PHOTO}), "text", lambda s: _qr_decode(s)),
        ActionSpec("txt", "📄 Текст TXT", frozenset({PDF, SUB, DOCX}), "document",
                   lambda s: asyncio.to_thread(_to_txt, s)),
        ActionSpec("pdfsplit", "✂️ По страницам (ZIP)", frozenset({PDF}), "document",
                   lambda s: asyncio.to_thread(_pdf_split_sync, s)),
        ActionSpec("pdfcomp", "📦 Сжать PDF", frozenset({PDF}), "document",
                   lambda s: _pdf_compress(s)),
        ActionSpec("pdfmake", "📄 Собрать PDF", frozenset({TEXT}), "document",
                   lambda s: asyncio.to_thread(_txt_to_pdf_sync, s)),
        ActionSpec("qrmake", "🔗 Сделать QR", frozenset({TEXT}), "document",
                   lambda s: asyncio.to_thread(_qr_from_text_sync, s)),
        ActionSpec("csv", "📊 CSV", frozenset({XLSX}), "document",
                   lambda s: asyncio.to_thread(_xlsx_to_csv_sync, s)),
        ActionSpec("xlsx", "📊 XLSX", frozenset({CSV}), "document",
                   lambda s: asyncio.to_thread(_csv_to_xlsx_sync, s)),
        ActionSpec("info", "ℹ️ Инфо о файле", frozenset({VIDEO, VIDEO_NOTE, AUDIO, VOICE, PHOTO, PDF, DOCX, XLSX, CSV, TEXT, SUB}), "text",
                   lambda s: _media_info(s)),
        ActionSpec("hash", "#️⃣ Хеши MD5/SHA256", frozenset({VIDEO, VIDEO_NOTE, AUDIO, VOICE, PHOTO, PDF, DOCX, XLSX, CSV, TEXT, SUB}), "text",
                   lambda s: _hash_file(s)),
        ActionSpec("unzip", "📤 Распаковать архив", frozenset({ZIP}), "document",
                   lambda s, k: asyncio.to_thread(
                       _unzip_files, s,
                       s.with_name(s.stem + "_unpacked"),
                   )),
    )
}


def _to_txt(src: Path) -> Path:
    ext = src.suffix.lower()
    out = _out(src, ".txt")
    if ext == ".pdf":
        reader = PdfReader(str(src))
        text = "\n\n".join((page.extract_text() or "") for page in reader.pages)
    elif ext == ".docx":
        text = _docx_to_plain(src)
    else:
        return _sub_to_txt_sync(src)
    out.write_text(text, encoding="utf-8")
    return out


def _docx_to_plain(src: Path) -> str:
    from docx import Document

    doc = Document(str(src))
    return "\n".join(p.text for p in doc.paragraphs)


def options_for(kind: str) -> list[ActionSpec]:
    return [spec for spec in ACTIONS.values() if kind in spec.kinds]


_GROUPS = {
    "voice": "🎧 Извлечь звук",
    "mp3": "🎧 Извлечь звук",
    "wav": "🎧 Извлечь звук",
    "m4a": "🎧 Извлечь звук",
    "flac": "🎧 Извлечь звук",
    "ogg": "🎧 Извлечь звук",
    "ringtone": "🎧 Извлечь звук",
    "gif": "🎞 Формат видео",
    "mp4": "🎞 Формат видео",
    "webm": "🎞 Формат видео",
    "mute": "🎞 Формат видео",
    "vsticker": "🎞 Формат видео",
    "speed15": "🎚 Темп",
    "speed2": "🎚 Темп",
    "unzip": "📤 Распаковка",
    "jpg": "🖼 Конвертировать в",
    "png": "🖼 Конвертировать в",
    "webp": "🖼 Конвертировать в",
    "bmp": "🖼 Конвертировать в",
    "tiff": "🖼 Конвертировать в",
    "ico": "🖼 Конвертировать в",
    "compress": "📦 Оптимизация",
    "sticker": "📦 Оптимизация",
    "gray": "✨ Эффекты",
    "invert": "✨ Эффекты",
    "rotate": "✨ Эффекты",
    "txt": "📄 Документы",
    "pdfsplit": "✂️ PDF-инструменты",
    "pdfcomp": "✂️ PDF-инструменты",
    "pdfmake": "📄 Документы",
    "qrmake": "🔗 QR-коды",
    "csv": "📊 Таблицы",
    "xlsx": "📊 Таблицы",
    "qrscan": "🔍 QR-коды",
    "info": "🔧 Инструменты",
    "hash": "🔧 Инструменты",
}

ACTIONS = {
    code: replace(spec, group=_GROUPS.get(code, ""))
    for code, spec in ACTIONS.items()
}
